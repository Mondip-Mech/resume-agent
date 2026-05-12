"""
agents/application_agent.py
────────────────────────────
Assembles the final application package:
  • Tailored cover letter
  • Formatted resume (DOCX + plain text)
  • Submission readiness check
  • Optional: API submission stub (LinkedIn, email)
"""
from __future__ import annotations

import logging
from pathlib import Path

from core.llm_client import LLMClient
from core.models import (
    AnalysisReport,
    ApplicationPackage,
    CoverLetter,
    ExtractedJD,
    ExtractedResume,
    RewrittenResume,
)

logger = logging.getLogger(__name__)


class ApplicationAgent:
    """
    Persona: Expert career consultant.
    Outputs: Cover letter + formatted resume package.
    Safety: Never handles live credentials or actual submission without human review.
    """

    SYSTEM_PROMPT = """You are an expert career consultant writing a cover letter for a job application.

COVER LETTER RULES:
- Opening hook: Lead with a specific, quantifiable achievement relevant to the role
- Body (2 paragraphs): Bridge top 3 candidate strengths to top 3 role needs
- Closing: Clear call to action ("I'd welcome a conversation...")
- Length: Exactly 3 paragraphs, 250-300 words total
- Tone: Warm, confident, human — never sycophantic or robotic
- Mirror JD language without copy-pasting it
- Reference company by name (not "your company")
- Do NOT start with "I am writing to apply for..."
- Do NOT use "I am passionate about..." or similar clichés"""

    def __init__(self, llm: LLMClient, output_dir: str = "./output"):
        self.llm = llm
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)

    def assemble(
        self,
        jd: ExtractedJD,
        resume: ExtractedResume,
        rewritten: RewrittenResume,
        analysis: AnalysisReport,
        session_id: str,
    ) -> ApplicationPackage:
        """Generate cover letter and assemble the full application package."""
        logger.info("ApplicationAgent: Assembling application package...")

        # 1. Generate cover letter (isolated try/except — failure here must not
        #    kill the resume DOCX which is already written at this point)
        try:
            cover_letter = self._write_cover_letter(jd, resume, rewritten, analysis)
        except Exception as cl_err:
            logger.warning(f"Cover letter generation failed: {cl_err}. Using placeholder.")
            cover_letter = CoverLetter(
                body=(
                    f"Dear Hiring Team at {jd.company},\n\n"
                    "Please find attached my tailored resume for the "
                    f"{jd.job_title} role. I am very excited about this opportunity "
                    f"and confident my background aligns well with your requirements.\n\n"
                    "I would welcome a conversation to discuss how I can contribute "
                    f"to {jd.company}'s mission.\n\n"
                    "Best regards,\n"
                    f"{resume.name}"
                ),
                subject_line=f"Re: {jd.job_title} — {resume.name}",
                tone="professional",
            )

        # 2. Format resume as DOCX
        resume_docx_path = self._save_resume_docx(rewritten, session_id)
        cover_letter_docx_path = self._save_cover_letter_docx(cover_letter, jd, resume, session_id)

        # 3. Submission readiness check
        ready, notes = self._readiness_check(rewritten, analysis)

        return ApplicationPackage(
            job_id=session_id,
            cover_letter=cover_letter,
            resume_text=rewritten.full_text,
            resume_docx_path=str(resume_docx_path) if resume_docx_path else None,
            cover_letter_docx_path=str(cover_letter_docx_path) if cover_letter_docx_path else None,
            ats_score_before=analysis.ats_score.overall,
            ats_score_after=rewritten.new_ats_score.overall if rewritten.new_ats_score else None,
            keywords_added=rewritten.added_keywords,
            ready_to_submit=ready,
            submission_notes=notes,
        )

    # ─── Cover Letter ──────────────────────────────────────────

    def _write_cover_letter(
        self,
        jd: ExtractedJD,
        resume: ExtractedResume,
        rewritten: RewrittenResume,
        analysis: AnalysisReport,
    ) -> CoverLetter:
        """
        Generate a personalized cover letter.

        Uses a plain-text call (no JSON) so the body can never be truncated
        mid-sentence inside a JSON string value.
        """
        top_strengths = "\n".join(f"- {s}" for s in analysis.strengths[:4])
        key_requirements = "\n".join([
            f"- {r.text}"
            for r in jd.requirements[:4]
            if r.importance == "required"
        ])

        prompt = f"""Write a professional cover letter for this job application.

ROLE: {jd.job_title} at {jd.company}
CANDIDATE: {resume.name}, currently {resume.current_title}

TOP CANDIDATE STRENGTHS:
{top_strengths}

KEY JD REQUIREMENTS TO ADDRESS:
{key_requirements}

CANDIDATE SUMMARY:
{resume.summary[:400]}

CULTURE / TONE SIGNALS FROM JD:
{', '.join(jd.culture_signals[:5])}

COVER LETTER RULES:
- Exactly 3 paragraphs, 250-300 words total
- Opening: lead with a specific achievement relevant to the role
- Body: bridge top 3 candidate strengths to top 3 role needs
- Closing: clear call to action ("I would welcome a conversation...")
- Warm, confident tone — not sycophantic or robotic
- Reference {jd.company} by name
- Do NOT start with "I am writing to apply for..."
- Do NOT use "I am passionate about..."

OUTPUT INSTRUCTIONS:
- Return ONLY the cover letter body as plain text.
- Do NOT include a subject line, greeting, or sign-off.
- Do NOT wrap in JSON, markdown, or code fences.
- Start directly with the first sentence of the first paragraph."""

        body = self.llm.call(prompt, system=self.SYSTEM_PROMPT, max_tokens=600)

        # Clean up any accidental JSON wrapping the model might add
        body = body.strip()
        if body.startswith("{") and "body" in body:
            try:
                import json as _json
                parsed = _json.loads(body)
                body = parsed.get("body", body)
            except Exception:
                pass

        subject_line = f"Re: {jd.job_title} — {resume.name}"

        return CoverLetter(
            body=body,
            subject_line=subject_line,
            tone="professional",
        )

    # ─── Document Generation ──────────────────────────────────

    def _save_resume_docx(self, rewritten: RewrittenResume, session_id: str) -> Path | None:
        """Save resume as ATS-safe DOCX."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt

            doc = Document()

            # Set margins
            for section in doc.sections:
                section.top_margin = Inches(0.75)
                section.bottom_margin = Inches(0.75)
                section.left_margin = Inches(1.0)
                section.right_margin = Inches(1.0)

            # Write content (simple line-by-line, ATS safe)
            for line in rewritten.full_text.split("\n"):
                line = line.strip()
                if not line:
                    doc.add_paragraph()
                    continue

                # Detect section headers (short, all-caps or title case lines)
                is_header = (
                    line.isupper() and len(line) < 40
                    or line.endswith(":") and len(line) < 40
                )

                para = doc.add_paragraph()
                run = para.add_run(line)

                if is_header:
                    run.bold = True
                    run.font.size = Pt(12)
                    para.paragraph_format.space_before = Pt(8)
                else:
                    run.font.size = Pt(10.5)

            path = self.output_dir / f"{session_id}_resume.docx"
            doc.save(str(path))
            logger.info(f"Resume saved: {path}")
            return path

        except ImportError:
            logger.warning("python-docx not installed. Saving as .txt instead.")
            path = self.output_dir / f"{session_id}_resume.txt"
            path.write_text(rewritten.full_text)
            return path
        except Exception as e:
            logger.error(f"Failed to save resume DOCX: {e}")
            return None

    def _save_cover_letter_docx(
        self, cover: CoverLetter, jd: ExtractedJD, resume: ExtractedResume, session_id: str
    ) -> Path | None:
        """Save cover letter as DOCX."""
        try:
            from docx import Document
            from docx.shared import Inches, Pt

            doc = Document()
            for section in doc.sections:
                section.top_margin = Inches(1.0)
                section.bottom_margin = Inches(1.0)
                section.left_margin = Inches(1.25)
                section.right_margin = Inches(1.25)

            # Header
            header = doc.add_paragraph()
            header.add_run(resume.name or "").bold = True
            doc.add_paragraph(resume.email or "")
            doc.add_paragraph()

            # Subject
            subj = doc.add_paragraph()
            subj.add_run(cover.subject_line).bold = True
            doc.add_paragraph()

            # Body (split by double newlines)
            for para_text in cover.body.split("\n\n"):
                if para_text.strip():
                    p = doc.add_paragraph()
                    p.add_run(para_text.strip())
                    p.paragraph_format.space_after = Pt(12)

            path = self.output_dir / f"{session_id}_cover_letter.docx"
            doc.save(str(path))
            logger.info(f"Cover letter saved: {path}")
            return path

        except Exception as e:
            logger.warning(f"DOCX save failed: {e}. Saving .txt")
            path = self.output_dir / f"{session_id}_cover_letter.txt"
            path.write_text(f"{cover.subject_line}\n\n{cover.body}")
            return path

    # ─── Readiness Check ──────────────────────────────────────

    def _readiness_check(self, rewritten: RewrittenResume, analysis: AnalysisReport):
        """Determine if the package is ready for submission."""
        issues = []
        ats = rewritten.new_ats_score

        if ats and ats.overall < 60:
            issues.append(f"ATS score is low ({ats.overall}/100). Consider manual review.")
        if ats and ats.keyword_match < 50:
            issues.append(f"Only {ats.keyword_match:.0f}% of JD keywords matched.")
        if ats and ats.quantification < 30:
            issues.append("Less than 30% of bullets have metrics. Consider adding numbers.")

        missing_required = [
            g.requirement for g in analysis.skill_gaps
            if g.match.value == "missing" and g.importance == "required"
        ]
        if missing_required:
            issues.append(
                f"Required skills still missing: {', '.join(missing_required[:3])}"
            )

        notes = (
            "\n".join(issues) if issues
            else "Package looks good. Human review recommended before submission."
        )
        ready = len(issues) == 0

        return ready, notes
